"""
writing_workflow/formatter.py
──────────────────────────────
Renders completed draft + rich elements to the final output format.
Each mode outputs a document that looks native to its format:

  thesis    → University thesis with front page, APA 7 formatting
  article   → IEEE/APA journal article with abstract box and keywords
  blog      → Clean web article with byline and pull-quote styling
  tech_doc  → Technical doc with version header and monospace code
  report    → Business report with executive summary block
"""

from __future__ import annotations
import os, re, json, textwrap
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime

from config.modes import OutputModeConfig
from writing_workflow.outline_gen import Outline
from writing_workflow.section_drafter import DraftedSection
from writing_workflow.rich_content_generator import RichElement
from writing_workflow.document_templates import DocumentMeta, get_template
from writing_workflow.journal_styles import get_journal_style


@dataclass
class FormattedOutput:
    file_path:     str
    format:        str
    word_count:    int
    section_count: int
    figure_count:  int


def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")[:60]


def _today() -> str:
    return datetime.now().strftime("%B %d, %Y")


def _label_elements(sections: list[DraftedSection]) -> None:
    fig_n = table_n = eq_n = code_n = 1
    for sec in sections:
        for el in sec.rich_elements:
            if el.type in ("mermaid", "chart"):
                el.label = f"Figure {fig_n}"; fig_n += 1
            elif el.type == "table":
                el.label = f"Table {table_n}"; table_n += 1
            elif el.type == "latex":
                el.label = f"Equation {eq_n}"; eq_n += 1
            elif el.type == "code":
                el.label = f"Listing {code_n}"; code_n += 1


# ─────────────────────────────────────────────────────────────────────────────
# DOCX builders
# ─────────────────────────────────────────────────────────────────────────────

def _apply_template_to_doc(doc, template, meta: DocumentMeta):
    """Apply margins, fonts, and page setup from template."""
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    section = doc.sections[0]
    section.top_margin    = Cm(template.margin_top_cm)
    section.bottom_margin = Cm(template.margin_bottom_cm)
    section.left_margin   = Cm(template.margin_left_cm)
    section.right_margin  = Cm(template.margin_right_cm)

    # Default paragraph style — font + line spacing
    style = doc.styles["Normal"]
    style.font.name = template.font_name
    style.font.size = Pt(template.font_size_pt)
    from docx.shared import Pt as Pt2
    from docx.oxml.ns import qn as qn2
    pPr = style.paragraph_format
    pPr.line_spacing = template.line_spacing * 12 * 914  # line spacing in twips approx
    pPr.space_after  = Pt2(0)

    # Page numbers in footer
    if template.page_numbers and template.page_number_position == "bottom_center":
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = 1  # CENTER
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def _thesis_front_page(doc, title: str, meta: DocumentMeta):
    """APA 7 thesis title page."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph("\n" * 4)  # push content down ~1/3 page

    # Title — bold, centred, title case, max 12 words on one line
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold      = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(meta.author or "Author Name")

    # Affiliation block
    doc.add_paragraph()
    for line in [meta.department, meta.university]:
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)

    doc.add_paragraph()

    # Degree + submission statement
    if meta.degree:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stmt = meta.get_submission_statement()
        p.add_run(stmt)

    doc.add_paragraph()

    # Supervisor
    if meta.supervisor:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Supervisor: {meta.supervisor}")

    # Student ID
    if meta.student_id:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Student ID: {meta.student_id}")

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(meta.get_submission_date())

    doc.add_page_break()


def _article_title_block(doc, title: str, meta: DocumentMeta, abstract_text: str, keywords: list[str], style=None):
    """IEEE/APA article title block with abstract box and keywords."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold      = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Authors + affiliation
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(meta.author)
    if meta.affiliation:
        doc.add_paragraph(meta.affiliation).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if meta.email:
        doc.add_paragraph(meta.email).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Abstract box
    if abstract_text:
        abs_para = doc.add_paragraph()
        style_name = style.name if style else "nature"
        abs_label  = "Abstract" if style_name in ("nature", "acm") else "Abstract"
        abs_run    = abs_para.add_run(f"{abs_label}  ")
        abs_run.bold = True
        abs_para.add_run(abstract_text[:600].replace("\n", " "))
        abs_para.paragraph_format.left_indent  = Pt(36)
        abs_para.paragraph_format.right_indent = Pt(36)

    # Keywords
    if keywords or meta.keywords:
        kw = keywords or meta.keywords
        kw_para = doc.add_paragraph()
        kw_para.paragraph_format.left_indent = Pt(36)
        kw_label = style.keywords_label if style else "Keywords:"
        bold_run = kw_para.add_run(f"{kw_label} ")
        bold_run.bold = True
        kw_para.add_run(", ".join(kw))

    doc.add_paragraph()


def _report_cover(doc, title: str, meta: DocumentMeta):
    """Business report cover page."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph("\n" * 6)

    # Organisation line
    if meta.organisation:
        org = doc.add_paragraph()
        org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = org.add_run(meta.organisation.upper())
        r.font.size = Pt(10)

    doc.add_paragraph()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(20)

    # Tagline
    if meta.tagline:
        tl = doc.add_paragraph()
        tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tl.add_run(meta.tagline)

    doc.add_paragraph("\n" * 8)

    # Author + date at bottom
    footer_block = doc.add_paragraph()
    footer_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_block.add_run(f"{meta.author}  |  {meta.date or _today()}")

    doc.add_page_break()


def _tech_doc_header(doc, title: str, meta: DocumentMeta):
    """Technical doc header with version info."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    info = doc.add_paragraph()
    info.add_run(f"Author: {meta.author}    Date: {meta.date or _today()}")
    info.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Horizontal rule
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p     = doc.add_paragraph()
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()


def _rich_to_docx(el: RichElement, doc) -> None:
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if el.type in ("mermaid", "chart", "latex") and el.image_path and Path(el.image_path).exists():
        doc.add_paragraph()
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(el.image_path, width=Inches(5.0))
        cap = doc.add_paragraph(f"{el.label}: {el.caption}" if el.label else el.caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic    = True
        cap.runs[0].bold      = True

    elif el.type == "table":
        rows = [r for r in el.source.strip().splitlines() if r.strip().startswith("|")]
        rows = [r for r in rows if not re.match(r"^\|[-:| ]+\|$", r.strip())]
        if rows:
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            ncols = max(len(row) for row in cells)
            table = doc.add_table(rows=len(cells), cols=ncols)
            table.style = "Table Grid"
            for i, row in enumerate(cells):
                for j, cell_text in enumerate(row[:ncols]):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    if i == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            cap = doc.add_paragraph(f"{el.label}: {el.caption}" if el.label else el.caption)
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic    = True
            cap.runs[0].bold      = True

    elif el.type == "code":
        code_text = re.sub(r"^```\w*\n?", "", el.source).rstrip("`").strip()
        if el.label:
            lbl = doc.add_paragraph(f"{el.label}: {el.caption}")
            lbl.runs[0].font.size = Pt(9)
            lbl.runs[0].bold      = True
            lbl.runs[0].italic    = True
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def _build_docx(title, sections, mode, meta, output_path):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc      = Document()
    if mode.name == "article":
        jstyle   = get_journal_style(meta.journal_style)
        template = get_template(mode.name)
        # Override template with journal style settings
        template.line_spacing  = jstyle.line_spacing
        template.font_name     = jstyle.font_name
        template.font_size_pt  = jstyle.font_size_pt
    else:
        template = get_template(mode.name)
    _apply_template_to_doc(doc, template, meta)

    # Extract abstract text for article title block
    abstract_section = next((s for s in sections if s.key == "abstract"), None)
    abstract_text    = abstract_section.content[:600] if abstract_section else ""

    # ── Mode-specific front matter ─────────────────────────────────────────
    if mode.name == "thesis":
        _thesis_front_page(doc, title, meta)
    elif mode.name == "article":
        jstyle = get_journal_style(meta.journal_style)
        _article_title_block(doc, title, meta, abstract_text,
                             meta.keywords or [], style=jstyle)
    elif mode.name == "report":
        _report_cover(doc, title, meta)
    elif mode.name == "tech_doc":
        _tech_doc_header(doc, title, meta)
    else:
        # Blog: simple byline
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        byline = doc.add_paragraph(f"By {meta.author}  |  {meta.date or _today()}")
        byline.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ── Table of Contents (thesis only) ───────────────────────────────────
    if mode.name in ("thesis", "article"):
        toc_sec = next((s for s in sections if s.key == "table_of_contents"), None)
        if toc_sec:
            doc.add_page_break()
            toc_h = doc.add_heading("Table of Contents", level=1)
            toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in toc_h.runs:
                run.bold = False
            doc.add_paragraph()
            for s in sections:
                if s.key not in ("title_page", "table_of_contents", "abstract", "references"):
                    p   = doc.add_paragraph(style="No Spacing")
                    run = p.add_run(s.title)
                    run.font.size = Pt(12)
            doc.add_page_break()

    # ── Body sections ──────────────────────────────────────────────────────
    for sec in sections:
        # Skip sections already handled in front matter
        if sec.key in ("title_page", "table_of_contents"):
            continue
        # For article: abstract already in title block — skip re-printing
        if mode.name == "article" and sec.key == "abstract":
            continue

        # Section heading — use journal style names for articles
        if mode.name == "article":
            jstyle   = get_journal_style(meta.journal_style)
            sec_name = jstyle.section_names.get(sec.key, sec.title)
        else:
            sec_name = sec.title

        if sec.key == "references":
            ref_h = doc.add_heading("References", level=1)
            if mode.name in ("thesis", "article"):
                ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_heading(sec_name, level=1)

        # Body paragraphs
        for para_text in sec.content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("## "):
                doc.add_heading(para_text[3:], level=2)
            elif para_text.startswith("### "):
                doc.add_heading(para_text[4:], level=3)
            elif para_text.startswith(("- ", "* ")):
                for item in para_text.split("\n"):
                    item = item.lstrip("-* ").strip()
                    if item:
                        doc.add_paragraph(item, style="List Bullet")
            else:
                doc.add_paragraph(para_text)

        # Rich elements
        for el in sec.rich_elements:
            _rich_to_docx(el, doc)

        # Page break after each major chapter in thesis
        if mode.name == "thesis" and sec.key not in ("abstract", "references", "conclusion"):
            pass  # don't force page break mid-document — let word wrap naturally

    doc.save(str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# Markdown builders
# ─────────────────────────────────────────────────────────────────────────────

def _rich_to_markdown(el: RichElement) -> str:
    label_line = f"*{el.label}: {el.caption}*\n\n" if el.label else ""
    if el.type == "mermaid":
        return f"\n\n{label_line}```mermaid\n{el.source}\n```\n"
    elif el.type == "table":
        return f"\n\n{label_line}{el.source}\n"
    elif el.type == "latex":
        return f"\n\n{label_line}$$\n{el.source}\n$$\n"
    elif el.type == "chart":
        if el.image_path and Path(el.image_path).exists():
            return f"\n\n{label_line}![{el.caption}]({el.image_path})\n"
        return f"\n\n*[Chart: {el.caption}]*\n"
    elif el.type == "code":
        return f"\n\n{label_line}{el.source}\n"
    return ""


def _build_markdown(title, sections, mode, meta):
    lines = [f"# {title}", ""]

    # Mode-specific header block
    if mode.name == "thesis":
        for line in [meta.author, meta.degree, meta.department,
                     meta.university, f"Supervisor: {meta.supervisor}" if meta.supervisor else "",
                     meta.date or _today()]:
            if line:
                lines += [f"*{line}*"]
        lines += ["", "---", ""]

    elif mode.name == "article":
        lines += [f"**{meta.author}**"]
        if meta.affiliation:
            lines.append(f"*{meta.affiliation}*")
        if meta.keywords:
            lines.append(f"**Keywords:** {', '.join(meta.keywords)}")
        lines += ["", "---", ""]

    elif mode.name == "blog":
        lines += [f"*By {meta.author}  |  {meta.date or _today()}*", "", "---", ""]

    elif mode.name == "report":
        lines += [f"**{meta.organisation}**" if meta.organisation else "",
                  f"*{meta.author}  |  {meta.date or _today()}*", "", "---", ""]

    elif mode.name == "tech_doc":
        lines += [f"**Author:** {meta.author}  |  **Date:** {meta.date or _today()}",
                  "", "---", ""]

    for sec in sections:
        if sec.key in ("title_page", "table_of_contents"):
            continue
        lines += [f"## {sec.title}", "", sec.content]
        for el in sec.rich_elements:
            lines.append(_rich_to_markdown(el))
        lines.append("")

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# HTML builder
# ─────────────────────────────────────────────────────────────────────────────

def _rich_to_html(el: RichElement) -> str:
    label_html = f'<p class="caption"><strong>{el.label}:</strong> <em>{el.caption}</em></p>' if el.label else ""
    if el.type == "mermaid":
        return f'<div class="rich-block"><div class="mermaid">{el.source}</div>{label_html}</div>'
    elif el.type == "table":
        return f'<div class="rich-block">{label_html}<div class="md-table">{el.source}</div></div>'
    elif el.type == "latex":
        return f'<div class="rich-block equation"><div class="math">\\[{el.source}\\]</div>{label_html}</div>'
    elif el.type == "chart":
        try:
            spec = json.loads(el.source)
            uid  = abs(hash(el.source)) % 100000
            return textwrap.dedent(f"""
                <div class="rich-block">
                  <div id="chart_{uid}" style="width:100%;height:380px;"></div>
                  {label_html}
                  <script>(function(){{
                    var spec={json.dumps(spec)};
                    var traces=spec.data.map(function(s){{
                      return {{x:spec.x_labels,y:s.values,name:s.series,
                               type:spec.type==='line'?'scatter':spec.type,
                               mode:spec.type==='line'?'lines+markers':undefined}};
                    }});
                    Plotly.newPlot('chart_{uid}',traces,
                      {{title:spec.title,xaxis:{{title:spec.x_label}},
                        yaxis:{{title:spec.y_label}},margin:{{t:40}}}});
                  }})();</script>
                </div>""")
        except Exception:
            return f'<div class="rich-block"><p>[Chart: {el.caption}]</p></div>'
    elif el.type == "code":
        lang = "python"
        code = el.source
        m = re.match(r"^```(\w+)\n(.*?)```$", el.source, re.DOTALL)
        if m:
            lang, code = m.group(1), m.group(2)
        safe = code.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        return f'<div class="rich-block">{label_html}<pre><code class="language-{lang}">{safe}</code></pre></div>'
    return ""


def _build_html(title, sections, mode, meta):
    # Mode-specific CSS and header
    mode_styles = {
        "thesis":   "font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 2;",
        "article":  "font-family: 'Times New Roman', serif; font-size: 11pt; line-height: 1.5;",
        "blog":     "font-family: -apple-system, sans-serif; font-size: 18px; line-height: 1.8;",
        "tech_doc": "font-family: 'Segoe UI', sans-serif; font-size: 14px; line-height: 1.5;",
        "report":   "font-family: 'Calibri', sans-serif; font-size: 11pt; line-height: 1.4;",
    }
    body_style = mode_styles.get(mode.name, mode_styles["report"])

    # Build header block per mode
    if mode.name == "thesis":
        header_html = textwrap.dedent(f"""
            <div class="front-page" style="text-align:center;padding:4em 0;">
              <h1 style="font-size:1.6em;">{title}</h1>
              <p>{meta.author}</p>
              <p>{meta.degree}</p>
              <p>{meta.department}<br>{meta.university}</p>
              {'<p>Supervisor: ' + meta.supervisor + '</p>' if meta.supervisor else ''}
              <p>{meta.date or _today()}</p>
            </div><hr>""")
    elif mode.name == "article":
        kw = ", ".join(meta.keywords) if meta.keywords else ""
        header_html = textwrap.dedent(f"""
            <div class="article-header" style="text-align:center;margin-bottom:2em;">
              <h1>{title}</h1>
              <p><strong>{meta.author}</strong></p>
              {'<p><em>' + meta.affiliation + '</em></p>' if meta.affiliation else ''}
              {'<p><strong>Keywords:</strong> ' + kw + '</p>' if kw else ''}
            </div>""")
    elif mode.name == "blog":
        header_html = textwrap.dedent(f"""
            <div class="byline" style="margin-bottom:2em;">
              <h1>{title}</h1>
              <p style="color:#666;font-size:0.9em;">By <strong>{meta.author}</strong> &nbsp;|&nbsp; {meta.date or _today()}</p>
              <hr>
            </div>""")
    elif mode.name == "report":
        header_html = textwrap.dedent(f"""
            <div class="report-cover" style="border-bottom:3px solid #333;padding-bottom:1em;margin-bottom:2em;">
              {'<p style="font-size:0.8em;text-transform:uppercase;letter-spacing:2px;">' + meta.organisation + '</p>' if meta.organisation else ''}
              <h1>{title}</h1>
              <p style="color:#666;">{meta.author} &nbsp;|&nbsp; {meta.date or _today()}</p>
            </div>""")
    else:
        header_html = f"<h1>{title}</h1><p><em>{meta.author} | {meta.date or _today()}</em></p><hr>"

    body_parts = [header_html]
    for sec in sections:
        if sec.key in ("title_page", "table_of_contents"):
            continue
        safe = sec.content.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        paras = [f"<p>{p.strip()}</p>" for p in safe.split("\n\n") if p.strip()]
        body_parts.append(f'<section id="{sec.key}"><h2>{sec.title}</h2>')
        body_parts.extend(paras)
        for el in sec.rich_elements:
            body_parts.append(_rich_to_html(el))
        body_parts.append("</section>")

    return textwrap.dedent(f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>{title}</title>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/plotly.js/2.32.0/plotly.min.js"></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/mermaid/10.9.0/mermaid.min.js"></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/mathjax/3.2.2/es5/tex-mml-chtml.min.js"></script>
  <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github.min.css">
  <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/highlight.min.js"></script>
  <style>
    body {{ {body_style} max-width:860px; margin:60px auto; padding:0 20px; color:#1a1a1a; }}
    h1   {{ font-size:1.8em; margin-bottom:0.3em; }}
    h2   {{ font-size:1.3em; margin-top:2.5em; padding-bottom:0.3em; border-bottom:1px solid #ddd; }}
    p    {{ margin:1em 0; text-align:justify; }}
    .rich-block {{ margin:2em 0; padding:1em; background:#f8f8f8; border-left:3px solid #ccc; border-radius:4px; }}
    .caption {{ font-size:0.88em; color:#444; text-align:center; margin-top:0.5em; }}
    .equation {{ text-align:center; }}
    pre  {{ background:#f5f5f5; padding:1em; overflow-x:auto; border-radius:4px; font-size:0.85em; }}
    table {{ border-collapse:collapse; width:100%; margin:1em 0; }}
    th,td {{ border:1px solid #ccc; padding:8px 12px; text-align:left; font-size:0.9em; }}
    th {{ background:#f0f0f0; font-weight:bold; }}
    @media print {{ body {{ margin:0; }} .rich-block {{ break-inside:avoid; }} }}
  </style>
</head>
<body>
  {"".join(body_parts)}
  <script>
    mermaid.initialize({{startOnLoad:true,theme:'default'}});
    document.addEventListener('DOMContentLoaded',function(){{
      document.querySelectorAll('pre code').forEach(el=>hljs.highlightElement(el));
    }});
  </script>
</body>
</html>""").strip()


# ─────────────────────────────────────────────────────────────────────────────
# Public interface
# ─────────────────────────────────────────────────────────────────────────────

def render(
    outline:    Outline,
    sections:   list[DraftedSection],
    mode:       OutputModeConfig,
    output_dir: str = "./outputs",
    meta:       DocumentMeta = None,
) -> FormattedOutput:
    if meta is None:
        meta = DocumentMeta()

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    _label_elements(sections)

    slug      = _slugify(outline.title)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename  = f"{slug}_{timestamp}{mode.file_extension}"
    out_path  = out_dir / filename

    total_words = sum(s.word_count for s in sections)
    total_figs  = sum(len(s.rich_elements) for s in sections)

    if mode.output_format == "markdown":
        content = _build_markdown(outline.title, sections, mode, meta)
        out_path.write_text(content, encoding="utf-8")

    elif mode.output_format == "docx":
        _build_docx(outline.title, sections, mode, meta, out_path)

    elif mode.output_format == "html":
        content = _build_html(outline.title, sections, mode, meta)
        out_path.write_text(content, encoding="utf-8")

    elif mode.output_format == "pdf":
        html_content = _build_html(outline.title, sections, mode, meta)
        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(str(out_path))
        except ImportError:
            md_path = out_path.with_suffix(".md")
            _build_markdown(outline.title, sections, mode, meta)
            out_path = md_path
            print(f"[formatter] weasyprint not installed — saved as markdown")

    print(f"[formatter] Saved → {out_path} ({total_words:,} words, {total_figs} figures)")
    return FormattedOutput(file_path=str(out_path), format=mode.output_format,
                           word_count=total_words, section_count=len(sections),
                           figure_count=total_figs)
